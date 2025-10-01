"""
redact_pdf.py

Provides redact_pdf_with_hits(file_bytes_or_path, hits, preview_mode, black_box)
- preview_mode=True => draws semi-transparent colored overlays (non-destructive preview)
- preview_mode=False => applies destructive redaction using black boxes (default behavior)
"""

import io
from typing import List, Optional, Dict, Union
import fitz  # PyMuPDF

# We import CATEGORY_COLORS only for consistent preview colors
from .extract_text import CATEGORY_COLORS, Hit

def _hex_to_rgb_floats(hex_color: str):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) / 255.0 for i in (0, 2, 4))

def redact_pdf_with_hits(
    input_source: Union[str, bytes],
    hits: List[Hit],
    output_path: Optional[str] = None,
    preview_mode: bool = False,
    black_box: bool = True,
) -> bytes:
    """
    input_source: either raw PDF bytes or a file path
    hits: list of Hit objects (only hits with bbox for PDFs will be drawn/redacted)
    preview_mode: if True, draw semi-transparent colored rects (keeps text)
    preview_mode False + black_box True: destructive black rectangles
    """

    # Read/prepare the doc
    if isinstance(input_source, (bytes, bytearray)):
        doc = fitz.open(stream=input_source, filetype="pdf")
    else:
        doc = fitz.open(input_source)

    # Group hits by page
    page_map: Dict[int, List[Hit]] = {}
    for h in hits:
        page_map.setdefault(h.page, []).append(h)

    # Iterate pages and draw/add redaction annotations
    for page_num, phits in page_map.items():
        page = doc[page_num]
        # Add annotations first, apply once per page for destructive redaction
        for h in phits:
            if not h.bbox:
                continue
            x0, y0, x1, y1 = h.bbox
            # expand small amount to ensure full coverage (1.2pt padding)
            pad = 1.2
            rect = fitz.Rect(x0 - pad, y0 - pad, x1 + pad, y1 + pad)

            if preview_mode:
                color = CATEGORY_COLORS.get(h.category, "#000000")
                rgb = _hex_to_rgb_floats(color)
                # draw a semi-transparent filled rectangle
                page.draw_rect(rect, color=rgb, fill=(*rgb, 0.22), width=0.6)
            else:
                # destructive redaction
                fill = (0, 0, 0) if black_box else (1, 1, 1)
                page.add_redact_annot(rect, fill=fill)

        if not preview_mode:
            # apply redactions on page (applies all redact annots added)
            try:
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
            except Exception:
                # fallback: apply without restricting images
                page.apply_redactions()

    out = io.BytesIO()
    doc.save(out, garbage=4)
    doc.close()
    data = out.getvalue()
    if output_path:
        with open(output_path, "wb") as f:
            f.write(data)
    return data

def save_masked_file(file_bytes: bytes, ext: str, hits: List[Hit]) -> bytes:
    """
    For non-pdf files: destructive masking by replacing occurrences with block characters.
    Replaces longer hits first to avoid partial overlap replacements.
    """
    if not hits:
        return file_bytes

    if ext == ".txt":
        text = file_bytes.decode("utf-8", errors="ignore")
        for h in sorted(hits, key=lambda x: len(x.text), reverse=True):
            text = text.replace(h.text, "█" * len(h.text))
        return text.encode("utf-8")

    if ext == ".docx":
        try:
            from docx import Document
            from io import BytesIO
        except Exception:
            return file_bytes
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            para_text = para.text
            new_text = para_text
            for h in sorted(hits, key=lambda x: len(x.text), reverse=True):
                new_text = new_text.replace(h.text, "█" * len(h.text))
            if new_text != para_text:
                # wipe runs, replace with a single run - preserves some formatting but simplifies
                for r in para.runs:
                    r.text = ""
                para.add_run(new_text)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    return file_bytes
